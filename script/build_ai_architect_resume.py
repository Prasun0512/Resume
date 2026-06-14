from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Prasun_Kumar_AI_ML_Architect_Resume.docx"


BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(20, 33, 61)
GRAY = RGBColor(90, 90, 90)


def set_run(run, size=10.5, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    set_run(r, size=12, bold=True, color=BLUE)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    set_run(p.add_run(text), size=9.5)


def add_role(doc: Document, title: str, org: str, dates: str, bullets: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    set_run(p.add_run(f"{title} | {org}"), size=10.2, bold=True, color=DARK)
    p.add_run("\n")
    set_run(p.add_run(dates), size=9.3, color=GRAY)
    for bullet in bullets:
        add_bullet(doc, bullet)


def add_compact_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    set_run(p.add_run(f"{label}: "), size=9.5, bold=True, color=DARK)
    set_run(p.add_run(value), size=9.5)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9.8)
    styles["Normal"].paragraph_format.space_after = Pt(3)
    styles["List Bullet"].font.name = "Calibri"
    styles["List Bullet"].font.size = Pt(9.5)
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.25)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.13)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    set_run(title.add_run("PRASUN KUMAR"), size=18, bold=True, color=DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    set_run(
        subtitle.add_run(
            "AI/ML Technical Lead | GenAI Solution Architect | Agentic AI & RAG Specialist"
        ),
        size=10.5,
        bold=True,
        color=BLUE,
    )

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(6)
    set_run(
        contact.add_run(
            "Pune, India | cehprasunsinha@yahoo.com | LinkedIn: linkedin.com/in/prasun-kumar-1708 | GitHub: github.com/Prasun0512 | Portfolio: prasun0512.github.io/Resume/"
        ),
        size=8.8,
        color=GRAY,
    )

    add_heading(doc, "Professional Summary")
    summary = doc.add_paragraph()
    summary.paragraph_format.line_spacing = 1.05
    set_run(
        summary.add_run(
            "AI/ML Technical Lead and GenAI Solution Architect with 8.5+ years of software engineering experience, including full-time engineering work since July 2017. Designs and delivers enterprise AI systems across Agentic AI, LLM applications, RAG, document intelligence, OCR + LLM extraction, AI automation, and cloud-native platforms. Strong focus on production readiness: event-driven architecture, human-in-the-loop review, idempotency, retries, DLQs, evaluation gates, PII controls, observability, and cost-aware LLM engineering."
        ),
        size=9.7,
    )

    add_heading(doc, "Core Skills")
    add_compact_line(doc, "AI & GenAI", "LLMs, GPT, Llama, Gemma, Qwen, Mistral, IBM Granite, prompt engineering, structured extraction")
    add_compact_line(doc, "Agentic AI", "LangGraph, LangChain, multi-agent workflows, tool calling, function calling, memory, approval gates")
    add_compact_line(doc, "RAG", "vector search, semantic search, hybrid search, embeddings, chunking, metadata filtering, retrieval evaluation")
    add_compact_line(doc, "ML / DL", "Scikit-Learn, XGBoost, NLP, classification, recommendation systems, feature engineering, PyTorch, TensorFlow, Hugging Face")
    add_compact_line(doc, "MLOps / LLMOps", "MLflow, Docker, Kubernetes, CI/CD, model registry, experiment tracking, evaluation, monitoring")
    add_compact_line(doc, "Cloud / Backend", "Azure OpenAI, Azure Functions, Azure AI Search, Blob Storage, Service Bus, Azure Document Intelligence, Python, FastAPI, REST APIs, PostgreSQL, MySQL")

    add_heading(doc, "Professional Experience")
    add_role(
        doc,
        "Technical Lead, AI and ML",
        "Harbinger Group",
        "Aug 2024 - Present",
        [
            "Lead architecture and delivery for enterprise GenAI, Agentic AI, RAG, document intelligence, and AI automation initiatives across Azure-based workflows.",
            "Architected a GenAI email-to-case automation platform using Azure Functions, Service Bus, Blob Storage, Azure OpenAI, OCR/document intelligence, MySQL, Graph API, and case-management integrations.",
            "Improved document extraction accuracy by 60%+ through better OCR normalization, prompt/schema design, confidence scoring, validation, and human review loops.",
            "Designed RAG patterns with chunking, semantic/hybrid retrieval, metadata filtering, citation support, and quality evaluation to improve grounded knowledge discovery.",
            "Led behavior scoring and text classification work using NLP/LLM patterns, including Llama/Gemma/Qwen experimentation, QLoRA fine-tuning concepts, adapter lifecycle practices, and evaluation dashboards.",
            "Reduced model/storage footprint by 77 TB through lifecycle cleanup, governance, artifact review, and storage optimization practices.",
            "Mentor engineers, lead design reviews, manage stakeholder discussions, and define production-readiness checks for AI delivery teams.",
        ],
    )
    add_role(
        doc,
        "Senior Software Engineer",
        "Harbinger Group",
        "Apr 2021 - Jul 2024",
        [
            "Delivered enterprise software and AI/ML-adjacent systems across backend services, automation workflows, integrations, and data-driven applications.",
            "Built and integrated APIs, data-processing flows, automation utilities, and reporting components with emphasis on maintainability and reliability.",
            "Collaborated with product, QA, DevOps, and client stakeholders to convert business requirements into scalable technical implementations.",
            "Improved engineering delivery through code reviews, reusable patterns, troubleshooting support, and mentoring of junior engineers.",
        ],
    )
    add_role(
        doc,
        "Software Engineer",
        "Extentia Information Technology",
        "Nov 2019 - Apr 2021",
        [
            "Developed backend and integration solutions using Python/.NET-style enterprise engineering practices, REST APIs, database workflows, and cloud-aligned delivery patterns.",
            "Participated in requirement analysis, implementation planning, testing support, and production troubleshooting for client-facing systems.",
        ],
    )
    add_role(
        doc,
        "Associate IT Applications Specialist",
        "Symantec Software India Pvt Ltd",
        "Jul 2017 - Oct 2019",
        [
            "Supported enterprise application engineering, automation, data workflows, and operational reliability for internal business systems.",
            "Built early foundation in production support, stakeholder communication, incident analysis, and disciplined engineering execution.",
        ],
    )
    add_role(
        doc,
        "Master Level Intern",
        "Symantec Software India Pvt Ltd",
        "Jan 2017 - Jul 2017",
        [
            "Worked on enterprise application support and automation foundations before moving into full-time engineering.",
        ],
    )

    add_heading(doc, "Selected AI Architecture Projects")
    projects = [
        ("GenAI-Powered Email-to-Case Automation Platform", "Technical Lead / AI Solution Architect", "Email ingestion, OCR, LLM extraction, automated case creation, Azure OpenAI, Azure Functions, Service Bus, Blob Storage, MySQL, retry handling, idempotency, audit trail, and human validation."),
        ("Enterprise RAG Knowledge Platform", "RAG Architect", "Document ingestion, chunking, embeddings, hybrid retrieval, Azure AI Search patterns, retrieval APIs, citations, confidence thresholds, and evaluation metrics."),
        ("Agentic AI Multi-Agent Platform", "Agentic AI Specialist", "Planner, research, retrieval, validator, executor, tool routing, human approval, retry, memory, and audit-trace patterns using LangGraph-style orchestration."),
        ("AI Behavioral Scoring Platform", "Technical Lead", "NLP/text classification experimentation using BERT/RoBERTa/SentenceTransformers/XGBoost-style patterns, multilingual support considerations, threshold calibration, and evaluation practices."),
        ("MLflow MLOps Platform", "AI Platform / MLOps", "Experiment tracking, evaluation gates, registry decisions, deployment readiness, monitoring handoffs, model lifecycle controls, and production governance patterns."),
    ]
    for name, role, detail in projects:
        add_bullet(doc, f"{name} ({role}): {detail}")

    add_heading(doc, "Public GitHub Portfolio")
    repos = [
        "enterprise-ai-ml-case-studies - sanitized enterprise GenAI/RAG/document AI/agentic workflow case studies with runnable POCs",
        "agentic-ai-langgraph-workflows - multi-agent workflow with approval gates, retries, tools, and audit traces",
        "enterprise-rag-quality-lab - RAG chunking, retrieval metrics, grounding, citation coverage, and evaluation examples",
        "document-intelligence-llm-pipeline - OCR + LLM extraction pipeline with redaction, schema validation, confidence scoring, and review routing",
        "email-to-case-genai-automation - event-driven GenAI automation with idempotency, retries, DLQ, validation, and audit logs",
        "ai-resume-job-matcher - explainable resume/JD matching with skill extraction, weighted scoring, gaps, and responsible AI notes",
    ]
    for repo in repos:
        add_bullet(doc, repo)

    add_heading(doc, "Education and Certification")
    add_bullet(doc, "Post Graduate Program in Machine Learning & Artificial Intelligence - Credential.net: credential.net/20453200-981b-4985-9ddf-2a44c5ee66cf")
    add_bullet(doc, "Program coverage: machine learning, deep learning, NLP, reinforcement learning, MLOps, and production AI foundations.")

    add_heading(doc, "Awards and Recognition")
    add_bullet(doc, "3x Quarterly Team Award Winner; 2x Superstar Award Winner; 4x Technical Star Recognition.")
    add_bullet(doc, "Additional recognitions include We-NOT-Me, going-the-extra-mile, Thank You, FunStar, PEP Silver, and Hackathon Recognition - 2017.")
    add_bullet(doc, "Recognized between 2017 and 2026 for technical excellence, innovation, leadership, ownership, collaboration, and delivery impact.")

    add_heading(doc, "Thought Leadership")
    add_bullet(doc, "Shares practical perspectives on Agentic AI, AI agents in production, GenAI architecture, RAG systems, LLM reliability, AI governance, and emerging AI trends.")
    add_bullet(doc, "YouTube: youtube.com/@AIWizardry277")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
